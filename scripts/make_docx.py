"""
Convert DOCUMENTATION.md to a formatted DOCUMENTATION.docx file.
Run: venv/Scripts/python.exe scripts/make_docx.py
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

SRC  = "DOCUMENTATION.md"
DEST = "DOCUMENTATION.docx"

# ── Colours ───────────────────────────────────────────────────────────────────
C_HEADING1  = RGBColor(0x1A, 0x53, 0x76)   # dark blue
C_HEADING2  = RGBColor(0x2E, 0x86, 0xAB)   # medium blue
C_HEADING3  = RGBColor(0x17, 0xA5, 0x89)   # teal
C_CODE_BG   = RGBColor(0xF2, 0xF2, 0xF2)   # light grey
C_CODE_TEXT = RGBColor(0x1E, 0x1E, 0x1E)   # near black
C_TABLE_HDR = RGBColor(0x1A, 0x53, 0x76)   # dark blue (same as H1)
C_ACCENT    = RGBColor(0xE8, 0x4C, 0x3B)   # red accent


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    # RGBColor stores as (r, g, b) tuple elements
    r, g, b = rgb[0], rgb[1], rgb[2]
    shd.set(qn("w:fill"),  f"{r:02X}{g:02X}{b:02X}")
    tcPr.append(shd)


def add_horizontal_rule(doc):
    p    = doc.add_paragraph()
    pPr  = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "2E86AB")
    pBdr.append(bot)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(4)


def style_heading(para, level: int):
    run = para.runs[0] if para.runs else para.add_run(para.text)
    run.bold = True
    if level == 1:
        run.font.size  = Pt(20)
        run.font.color.rgb = C_HEADING1
    elif level == 2:
        run.font.size  = Pt(15)
        run.font.color.rgb = C_HEADING2
    else:
        run.font.size  = Pt(12)
        run.font.color.rgb = C_HEADING3
    para.paragraph_format.space_before = Pt(14)
    para.paragraph_format.space_after  = Pt(4)


def add_code_block(doc, lines: list[str]):
    for line in lines:
        p   = doc.add_paragraph()
        run = p.add_run(line if line else " ")
        run.font.name        = "Courier New"
        run.font.size        = Pt(9)
        run.font.color.rgb   = C_CODE_TEXT
        # light grey shading on paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F2F2F2")
        pPr.append(shd)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.left_indent  = Cm(0.5)


def add_table_from_md(doc, rows: list[str]):
    # Parse markdown table rows
    data = []
    for row in rows:
        if re.match(r"\s*\|[-| :]+\|\s*$", row):
            continue  # separator row
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        data.append(cells)

    if not data:
        return

    cols   = len(data[0])
    table  = doc.add_table(rows=len(data), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for r_idx, row_data in enumerate(data):
        row = table.rows[r_idx]
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = cell_text
            run  = cell.paragraphs[0].runs[0] if cell.paragraphs[0].runs else cell.paragraphs[0].add_run(cell_text)
            if r_idx == 0:
                run.bold           = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                set_cell_bg(cell, C_TABLE_HDR)
            run.font.size = Pt(9.5)
            cell.paragraphs[0].paragraph_format.space_before = Pt(2)
            cell.paragraphs[0].paragraph_format.space_after  = Pt(2)

    doc.add_paragraph()  # breathing room after table


def build_docx():
    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── Default body font ─────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Cover title ───────────────────────────────────────────────────────────
    title = doc.add_heading("Multimodal AI Clinical Assistant", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = C_HEADING1
    title.runs[0].font.size = Pt(26)

    sub = doc.add_paragraph("Full Technical Documentation")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.size = Pt(13)
    sub.runs[0].font.color.rgb = C_HEADING2
    sub.runs[0].italic = True
    doc.add_paragraph()

    add_horizontal_rule(doc)
    doc.add_paragraph()

    # ── Parse markdown ────────────────────────────────────────────────────────
    with open(SRC, encoding="utf-8") as f:
        lines = f.readlines()

    i = 0
    while i < len(lines):
        line = lines[i].rstrip("\n")

        # Skip the top-level title (already added as cover)
        if line.startswith("# Multimodal AI Clinical Assistant"):
            i += 1
            continue

        # Horizontal rule
        if line.strip() in ("---", "***", "___"):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
            style_heading(p, 3)
            i += 1
            continue
        if line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
            style_heading(p, 2)
            i += 1
            continue
        if line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)
            style_heading(p, 1)
            i += 1
            continue

        # Code block
        if line.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i].rstrip("\n"))
                i += 1
            add_code_block(doc, code_lines)
            i += 1  # skip closing ```
            continue

        # Markdown table
        if line.startswith("|"):
            table_rows = []
            while i < len(lines) and lines[i].startswith("|"):
                table_rows.append(lines[i].rstrip("\n"))
                i += 1
            add_table_from_md(doc, table_rows)
            continue

        # Bullet list
        if line.startswith("- ") or line.startswith("* "):
            text = line[2:].strip()
            p    = doc.add_paragraph(style="List Bullet")
            # Handle inline bold/italic
            _add_inline(p, text)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Numbered list
        m = re.match(r"^\d+\.\s+(.*)", line)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_inline(p, m.group(1))
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # Blockquote (>) — render as indented italic
        if line.startswith("> "):
            p   = doc.add_paragraph()
            run = p.add_run(line[2:])
            run.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            p.paragraph_format.left_indent  = Cm(1.0)
            p.paragraph_format.space_after  = Pt(4)
            i += 1
            continue

        # Blank line
        if line.strip() == "":
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        _add_inline(p, line)
        p.paragraph_format.space_after = Pt(4)
        i += 1

    doc.save(DEST)
    print(f"Saved: {DEST}")


def _add_inline(para, text: str):
    """Parse inline **bold**, *italic*, and `code` markers."""
    # Pattern matches bold, italic, or code spans
    pattern = re.compile(r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`)")
    last = 0
    for m in pattern.finditer(text):
        # Text before this match
        if m.start() > last:
            para.add_run(text[last:m.start()])
        full = m.group(0)
        if full.startswith("**"):
            run = para.add_run(m.group(2))
            run.bold = True
        elif full.startswith("*"):
            run = para.add_run(m.group(3))
            run.italic = True
        elif full.startswith("`"):
            run = para.add_run(m.group(4))
            run.font.name      = "Courier New"
            run.font.size      = Pt(9.5)
            run.font.color.rgb = C_ACCENT
        last = m.end()
    if last < len(text):
        para.add_run(text[last:])


if __name__ == "__main__":
    build_docx()
